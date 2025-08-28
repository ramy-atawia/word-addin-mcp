"""
Enhanced Document Analyzer Service

Real document parsing and analysis using Azure OpenAI for summarization,
metadata extraction, and content analysis.
"""

import os
import json
import logging
from pathlib import Path
from typing import Dict, Any, Optional, List, Union
from datetime import datetime
import magic
import PyPDF2
from docx import Document
from bs4 import BeautifulSoup
from .llm_client import LLMClient
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument

logger = logging.getLogger(__name__)


class DocumentAnalyzerService:
    """Enhanced document analyzer with real parsing and LLM integration."""
    
    def __init__(self, azure_openai_api_key: Optional[str] = None, 
                 azure_openai_endpoint: Optional[str] = None,
                 azure_openai_deployment: Optional[str] = None):
        """
        Initialize the document analyzer service.
        
        Args:
            azure_openai_api_key: Azure OpenAI API key
            azure_openai_endpoint: Azure OpenAI endpoint URL
            azure_openai_deployment: Azure OpenAI deployment name
        """
        self.supported_formats = {
            '.pdf': 'PDF Document',
            '.docx': 'Word Document',
            '.doc': 'Word Document (Legacy)',
            '.html': 'HTML Document',
            '.htm': 'HTML Document',
            '.txt': 'Text Document',
            '.md': 'Markdown Document',
            '.rtf': 'Rich Text Document'
        }
        
        # Initialize LLM client
        self.llm_client = LLMClient(
            azure_openai_api_key=azure_openai_api_key,
            azure_openai_endpoint=azure_openai_endpoint,
            azure_openai_deployment=azure_openai_deployment
        )
        self.llm_available = self.llm_client.is_available()
        
        # Initialize LangChain components
        if self.llm_available:
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=4000,
                chunk_overlap=200,
                length_function=len,
                separators=["\n\n", "\n", " ", ""]
            )
        
        # Performance thresholds
        self.max_file_size = 100 * 1024 * 1024  # 100MB
        self.max_pages = 500  # Maximum pages to process
        
    def analyze_document(self, file_path: str, analysis_type: str = "comprehensive",
                        max_summary_length: Optional[int] = None,
                        include_metadata: bool = True,
                        include_keywords: bool = True) -> Dict[str, Any]:
        """
        Analyze a document with comprehensive parsing and LLM analysis.
        
        Args:
            file_path: Path to the document
            analysis_type: Type of analysis (comprehensive, summary, metadata, keywords)
            max_summary_length: Maximum length of summary
            include_metadata: Whether to include metadata analysis
            include_keywords: Whether to include keyword extraction
            
        Returns:
            Dictionary containing analysis results
        """
        try:
            start_time = datetime.now()
            
            # Validate file
            if not self._validate_file(file_path):
                return self._create_error_result("File validation failed")
            
            # Parse document
            parsed_content = self._parse_document(file_path)
            if not parsed_content.get("success"):
                return parsed_content
            
            content = parsed_content["content"]
            file_info = parsed_content["file_info"]
            
            # Perform analysis based on type
            analysis_results = {
                "status": "success",
                "analysis_type": analysis_type,
                "file_info": file_info,
                "processing_time": (datetime.now() - start_time).total_seconds(),
                "timestamp": datetime.now().isoformat()
            }
            
            # Basic content analysis
            if include_metadata:
                analysis_results["metadata"] = self._extract_basic_metadata(content)
            
            # LLM-powered analysis
            if self.llm_available:
                if analysis_type in ["comprehensive", "summary"]:
                    summary_result = self._generate_llm_summary(content, max_summary_length)
                    if summary_result.get("success"):
                        analysis_results["summary"] = summary_result["summary"]
                        analysis_results["key_points"] = summary_result["key_points"]
                
                if analysis_type in ["comprehensive", "keywords"] and include_keywords:
                    keywords_result = self._extract_keywords_with_llm(content)
                    if keywords_result.get("success"):
                        analysis_results["keywords"] = keywords_result["keywords"]
                        analysis_results["topics"] = keywords_result["topics"]
                
                if analysis_type == "comprehensive":
                    analysis_results["readability"] = self._analyze_readability_with_llm(content)
                    analysis_results["sentiment"] = self._analyze_sentiment_with_llm(content)
            
            # Add content statistics
            analysis_results["content_stats"] = self._calculate_content_statistics(content)
            
            logger.info(f"Document analysis completed for {file_path} in {analysis_results['processing_time']:.2f}s")
            return analysis_results
            
        except Exception as e:
            logger.error(f"Error analyzing document {file_path}: {str(e)}")
            return self._create_error_result(f"Analysis failed: {str(e)}")
    
    def _validate_file(self, file_path: str) -> bool:
        """Validate file for processing."""
        try:
            if not os.path.exists(file_path):
                logger.error(f"File not found: {file_path}")
                return False
            
            if not os.path.isfile(file_path):
                logger.error(f"Path is not a file: {file_path}")
                return False
            
            file_size = os.path.getsize(file_path)
            if file_size > self.max_file_size:
                logger.error(f"File too large: {file_size} bytes (max: {self.max_file_size})")
                return False
            
            # Check file extension
            file_ext = Path(file_path).suffix.lower()
            if file_ext not in self.supported_formats:
                logger.error(f"Unsupported file format: {file_ext}")
                return False
            
            return True
            
        except Exception as e:
            logger.error(f"File validation error: {str(e)}")
            return False
    
    def _parse_document(self, file_path: str) -> Dict[str, Any]:
        """Parse document based on file type."""
        try:
            file_ext = Path(file_path).suffix.lower()
            file_info = self._get_file_info(file_path)
            
            if file_ext == '.pdf':
                content = self._parse_pdf(file_path)
            elif file_ext == '.docx':
                content = self._parse_docx(file_path)
            elif file_ext in ['.html', '.htm']:
                content = self._parse_html(file_path)
            elif file_ext in ['.txt', '.md', '.rtf']:
                content = self._parse_text(file_path)
            else:
                return self._create_error_result(f"Unsupported file format: {file_ext}")
            
            if not content:
                return self._create_error_result("Failed to extract content from document")
            
            return {
                "success": True,
                "content": content,
                "file_info": file_info
            }
            
        except Exception as e:
            logger.error(f"Document parsing error: {str(e)}")
            return self._create_error_result(f"Parsing failed: {str(e)}")
    
    def _parse_pdf(self, file_path: str) -> str:
        """Parse PDF document using PyPDF2."""
        try:
            content = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check page count
                if len(pdf_reader.pages) > self.max_pages:
                    logger.warning(f"PDF has {len(pdf_reader.pages)} pages, limiting to {self.max_pages}")
                    pages_to_process = self.max_pages
                else:
                    pages_to_process = len(pdf_reader.pages)
                
                for page_num in range(pages_to_process):
                    try:
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text.strip():
                            content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"Failed to extract page {page_num + 1}: {str(e)}")
                        continue
            
            return content.strip()
            
        except Exception as e:
            logger.error(f"PDF parsing error: {str(e)}")
            raise
    
    def _parse_docx(self, file_path: str) -> str:
        """Parse Word document using python-docx."""
        try:
            doc = Document(file_path)
            content = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            # Extract table content
            for table in doc.tables:
                content += "\n--- Table ---\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    content += row_text + "\n"
                content += "--- End Table ---\n"
            
            return content.strip()
            
        except Exception as e:
            logger.error(f"DOCX parsing error: {str(e)}")
            raise
    
    def _parse_html(self, file_path: str) -> str:
        """Parse HTML document using BeautifulSoup."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                
                # Extract text from body
                if soup.body:
                    content = soup.body.get_text()
                else:
                    content = soup.get_text()
                
                # Clean up whitespace
                lines = (line.strip() for line in content.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                content = ' '.join(chunk for chunk in chunks if chunk)
                
                return content
                
        except Exception as e:
            logger.error(f"HTML parsing error: {str(e)}")
            raise
    
    def _parse_text(self, file_path: str) -> str:
        """Parse text-based documents."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        return content
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try with error handling
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                return content
                
        except Exception as e:
            logger.error(f"Text parsing error: {str(e)}")
            raise
    
    def _get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get comprehensive file information."""
        try:
            stat = os.stat(file_path)
            file_ext = Path(file_path).suffix.lower()
            
            # Detect MIME type
            try:
                mime_type = magic.from_file(file_path, mime=True)
            except:
                mime_type = "unknown"
            
            return {
                "name": Path(file_path).name,
                "path": file_path,
                "size": stat.st_size,
                "created": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "permissions": oct(stat.st_mode)[-3:],
                "owner": stat.st_uid,
                "group": stat.st_gid,
                "extension": file_ext,
                "mime_type": mime_type,
                "format": self.supported_formats.get(file_ext, "Unknown")
            }
            
        except Exception as e:
            logger.error(f"Error getting file info: {str(e)}")
            return {"error": str(e)}
    
    def _extract_basic_metadata(self, content: str) -> Dict[str, Any]:
        """Extract basic metadata from content."""
        try:
            lines = content.split('\n')
            words = content.split()
            
            # Detect language (basic English detection)
            english_words = sum(1 for word in words if word.isalpha() and len(word) > 2)
            total_words = len(words) if words else 1
            language_confidence = min(english_words / total_words, 1.0)
            
            return {
                "content_length": len(content),
                "words": len(words),
                "lines": len(lines),
                "paragraphs": len([line for line in lines if line.strip()]),
                "avg_line_length": sum(len(line) for line in lines) / len(lines) if lines else 0,
                "avg_word_length": sum(len(word) for word in words) / len(words) if words else 0,
                "has_unicode": any(ord(char) > 127 for char in content),
                "language_confidence": language_confidence,
                "estimated_language": "English" if language_confidence > 0.7 else "Unknown"
            }
            
        except Exception as e:
            logger.error(f"Error extracting basic metadata: {str(e)}")
            return {"error": str(e)}
    
    def _generate_llm_summary(self, content: str, max_length: Optional[int] = None) -> Dict[str, Any]:
        """Generate LLM-powered summary using the LLM client."""
        try:
            if not self.llm_available:
                return {"success": False, "error": "LLM not available"}
            
            # Split content into chunks if it's very long
            if len(content) > 8000:
                chunks = self.text_splitter.split_text(content)
                # Use first few chunks for summary
                content_for_summary = "\n\n".join(chunks[:3])
            else:
                content_for_summary = content
            
            # Generate comprehensive summary
            summary_result = self.llm_client.summarize_text(
                text=content_for_summary,
                summary_type="detailed",
                max_length=max_length
            )
            
            if not summary_result.get("success"):
                return summary_result
            
            # Extract key points
            key_points_result = self.llm_client.extract_keywords(
                text=summary_result["summary"],
                max_keywords=5
            )
            
            if not key_points_result.get("success"):
                key_points = "Key points extraction failed"
            else:
                key_points = ", ".join(key_points_result["keywords"])
            
            return {
                "success": True,
                "summary": summary_result["summary"],
                "key_points": key_points,
                "usage": summary_result.get("usage", {})
            }
            
        except Exception as e:
            logger.error(f"Error generating LLM summary: {str(e)}")
            return {"success": False, "error": str(e)}
    
    def _extract_keywords_with_llm(self, content: str) -> Dict[str, Any]:
        """Extract keywords and topics using LLM."""
        try:
            if not self.llm_available:
                return {"success": False, "error": "LLM not available"}
            
            # Extract keywords
            keywords_result = self.llm_client.extract_keywords(
                text=content[:2000],
                max_keywords=10
            )
            
            if not keywords_result.get("success"):
                return keywords_result
            
            # Extract topics using a custom prompt
            topics_prompt = f"""
            Analyze the following text and identify the 5 main topics or themes.
            Return only the topics, separated by commas.
            
            Text: {content[:2000]}...
            
            Topics:
            """
            
            topics_result = self.llm_client.generate_text(
                prompt=topics_prompt,
                max_tokens=150,
                temperature=0.3,
                system_message="You are an expert at topic identification. Return only the topics, separated by commas."
            )
            
            if not topics_result.get("success"):
                topics = ["Topic extraction failed"]
            else:
                topics_text = topics_result["text"].strip()
                topics = [topic.strip() for topic in topics_text.split(",") if topic.strip()]
            
            return {
                "success": True,
                "keywords": keywords_result["keywords"],
                "topics": topics,
                "usage": {
                    "keywords": keywords_result.get("usage", {}),
                    "topics": topics_result.get("usage", {})
                }
            }
            
        except Exception as e:
            logger.error(f"Error extracting keywords with LLM: {str(e)}")
            return {"success": False, "error": str(e)}
    
    def _analyze_readability_with_llm(self, content: str) -> Dict[str, Any]:
        """Analyze readability using LLM."""
        try:
            if not self.llm_available:
                return {"success": False, "error": "LLM not available"}
            
            return self.llm_client.analyze_readability(content[:1500])
            
        except Exception as e:
            logger.error(f"Error analyzing readability with LLM: {str(e)}")
            return {"success": False, "error": str(e)}
    
    def _analyze_sentiment_with_llm(self, content: str) -> Dict[str, Any]:
        """Analyze sentiment using LLM."""
        try:
            if not self.llm_available:
                return {"success": False, "error": "LLM not available"}
            
            return self.llm_client.analyze_sentiment(content[:1500])
            
        except Exception as e:
            logger.error(f"Error analyzing sentiment with LLM: {str(e)}")
            return {"success": False, "error": str(e)}
    
    def _calculate_content_statistics(self, content: str) -> Dict[str, Any]:
        """Calculate comprehensive content statistics."""
        try:
            lines = content.split('\n')
            words = content.split()
            
            # Word frequency analysis
            word_freq = {}
            for word in words:
                clean_word = word.lower().strip('.,!?;:()[]{}"\'-')
                if clean_word and len(clean_word) > 2:
                    word_freq[clean_word] = word_freq.get(clean_word, 0) + 1
            
            # Top words
            top_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]
            
            return {
                "total_characters": len(content),
                "total_words": len(words),
                "total_lines": len(lines),
                "total_paragraphs": len([line for line in lines if line.strip()]),
                "unique_words": len(word_freq),
                "avg_word_length": sum(len(word) for word in words) / len(words) if words else 0,
                "avg_line_length": sum(len(line) for line in lines) / len(lines) if lines else 0,
                "top_words": top_words,
                "word_diversity": len(word_freq) / len(words) if words else 0
            }
            
        except Exception as e:
            logger.error(f"Error calculating content statistics: {str(e)}")
            return {"error": str(e)}
    
    def _create_error_result(self, error_message: str) -> Dict[str, Any]:
        """Create standardized error result."""
        return {
            "status": "error",
            "error_message": error_message,
            "timestamp": datetime.now().isoformat()
        }


# Global instance for easy access
document_analyzer_service = DocumentAnalyzerService()
